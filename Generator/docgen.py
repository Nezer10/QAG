from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt
from firebaseConfig import store

bucketDoc = store.bucket()

class Docgen:
    def __init__(self, university:str, college: str, department: str, date: str, duration: str, exam_name: str,  lecturer: str, body_text: str,):
        self.university= university
        self.college= college
        self.department = department
        self.exam_name = exam_name
        self.date = date
        self.duration = duration
        self.lecturer = lecturer
        self.body_text = body_text  # Store body text as an instance variable
        self.doc = Document()
    
    def set_font(self):
     # Set font for the entire document
     style = self.doc.styles['Normal']
     font = style.font
     font.name = 'Times New Roman'
     font.size = Pt(12)  # Set font size, you can adjust it as needed
     font.element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')  # Ensure correct font for all regions

    def create_header(self):
     # Create header
     header = self.doc.sections[0].header
     # Create a table with 1 row and 3 columns
     header_table = header.add_table(rows=1, cols=3, width=Inches(15))
     

     # Left cell: university, college, department
     left_cell = header_table.cell(0, 0)
     left_paragraph = left_cell.paragraphs[0]
     left_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
     left_run = left_paragraph.add_run(f"University: {self.university}\nCollege: {self.college}\nDepartment: {self.department}")
     left_run.bold = True
     left_run.font.name = 'Times New Roman'

     # Center cell for logo
     center_cell = header_table.cell(0, 1)
     center_paragraph = center_cell.paragraphs[0]
     center_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

     # Add the logo image to the center cell
     center_paragraph.add_run().add_picture("Assets/Al-Mustansiriya_University_logo.png", width=Inches(1), height=Inches(1))  # Adjust width and height as needed

     # Right cell: exam name, date, duration
     right_cell = header_table.cell(0, 2)
     right_paragraph = right_cell.paragraphs[0]
     right_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
     right_run = right_paragraph.add_run(f"Duration: {self.duration}\nDate: {self.date}\nExam: {self.exam_name}")
     right_run.bold = True
     right_run.font.name = 'Times New Roman'

    def create_footer(self):
        # Create footer
        footer = self.doc.sections[0].footer
        table = footer.add_table(rows=1, cols=3, width=Inches(12))

        # Left side: Chairman signature
        left_cell = table.cell(0, 0)
        left_paragraph = left_cell.paragraphs[0]
        left_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        left_run = left_paragraph.add_run(f"Lecturer: {self.lecturer}")
        left_run.bold = True
        left_run.italic = True
        left_run.font.name = 'Times New Roman'
        

        # Center: page number (bold and italic)
        center_cell = table.cell(0, 1)
        center_paragraph = center_cell.paragraphs[0]
        center_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Insert the dynamic page number with bold and italic formatting
        center_run = center_paragraph.add_run("Page ")
        center_run.bold = True
        center_run.italic = True  # Italic for the "Page" label
        center_run.font.name = 'Times New Roman'
        self._add_page_number_field(center_run)  # Insert page number dynamically

        # Right side: Lecturer's name
        right_cell = table.cell(0, 2)
        right_paragraph = right_cell.paragraphs[0]
        right_run = right_paragraph.add_run(f"Chairman Signature:       ")
        right_run.bold = True
        right_run.italic = True  # Italic
        right_run.font.name = 'Times New Roman'
        right_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    def _add_page_number_field(self, run):
        fldChar1 = OxmlElement('w:fldChar')  # Create new element
        fldChar1.set(qn('w:fldCharType'), 'begin')  # Set attribute

        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')  # Sets space preserving
        instrText.text = "PAGE"  # Field code to insert page number

        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')

        run._r.append(fldChar1)
        run._r.append(instrText)
        run._r.append(fldChar2)
    
    def create_body(self):
        # Add body content with Arabic support
        body_paragraph = self.doc.add_paragraph(self.body_text)
        body_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT  # Set text alignment to Right for Arabic
        body_run = body_paragraph.add_run()
        body_run.font.name = 'Times New Roman'  # Use Times New Roman or any font that supports Arabic
        body_run.font.size = Pt(12)

        # Add right-to-left text direction
        rtl = OxmlElement('w:bidi')
        body_paragraph._p.get_or_add_pPr().append(rtl)

    def save_docx(self, filename, folder):
        doc_output = f"{filename}.docx"
        self.doc.save(doc_output)
        folder_path = f"{folder}/"
        blobDoc = bucketDoc.blob(folder_path + doc_output)
        blobDoc.upload_from_filename(doc_output)
        blobDoc.make_public()
        print(f"DOCX file {doc_output} generated successfully!")
        return doc_output
    
def docx_generate(filename, university, college, department, date, duration, exam_name, lecturer, body_text, folder):
        doc = Docgen(university, college, department, date, duration, exam_name, lecturer, body_text)
        doc.set_font()
        doc.create_header()
        doc.create_footer()
        doc.create_body()

        doc.save_docx(filename, folder)
    